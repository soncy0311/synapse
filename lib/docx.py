"""DOCX 파일 파서."""

from pathlib import Path

from docx import Document


def parse(path: str | Path) -> str:
    """DOCX 파일에서 텍스트를 추출한다."""
    doc = Document(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n".join(parts)
